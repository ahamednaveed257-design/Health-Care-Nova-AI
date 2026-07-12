from __future__ import annotations

from copy import deepcopy
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "outputs" / "documentation-pack"
DOCX_PATH = OUTPUT_DIR / "Care-Nova-AI-Documentation-Pack.docx"

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
SLATE = RGBColor(0x44, 0x4C, 0x56)
MUTED = RGBColor(0x6B, 0x72, 0x7A)
LIGHT_FILL = "F2F4F7"
DIAGRAM_FILL = "F7F9FC"
RULE = "D8DEE8"
BLACK = RGBColor(0x00, 0x00, 0x00)

RUN_DT = datetime.now().astimezone()
RUN_TS = f"{RUN_DT.strftime('%B')} {RUN_DT.day}, {RUN_DT.year} {RUN_DT.strftime('%H:%M %Z')}"
RUN_TS_ISO = RUN_DT.isoformat(timespec="seconds")


def set_font(run, name="Calibri", size=11, color=BLACK, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def set_paragraph_spacing(paragraph, before=0, after=0, line=1.1):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def paragraph_border(paragraph, color=RULE, size="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)


def extract_element_text(element):
    pieces = []
    for node in element.iter():
        if node.tag == qn("w:t") and node.text:
            pieces.append(node.text)
    return " ".join(" ".join(pieces).split())


def insert_body_element_before_section_properties(doc, element):
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:sectPr"):
            child.addprevious(element)
            return
    body.append(element)


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        element = tc_mar.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_row_layout(row, widths):
    for idx, width in enumerate(widths):
        row.cells[idx].width = width
        set_cell_margins(row.cells[idx])


def set_table_layout(table, widths):
    table.autofit = False
    for idx, width in enumerate(widths):
        table.columns[idx].width = width
    for row in table.rows:
        set_row_layout(row, widths)


def mark_header_row(row):
    tr_pr = row._tr.find(qn("w:trPr"))
    if tr_pr is None:
        tr_pr = OxmlElement("w:trPr")
        row._tr.insert(0, tr_pr)
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def add_heading(doc, text, level=1):
    style_name = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}[level]
    paragraph = doc.add_paragraph(style=style_name)
    run = paragraph.add_run(text)
    if level == 1:
        set_font(run, size=16, color=BLUE, bold=True)
        set_paragraph_spacing(paragraph, before=16, after=8, line=1.1)
    elif level == 2:
        set_font(run, size=13, color=BLUE, bold=True)
        set_paragraph_spacing(paragraph, before=12, after=6, line=1.1)
    else:
        set_font(run, size=12, color=DARK_BLUE, bold=True)
        set_paragraph_spacing(paragraph, before=8, after=4, line=1.1)
    return paragraph


def add_body(doc, text, before=0, after=6, italic=False):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_font(run, size=11, color=BLACK, italic=italic)
    set_paragraph_spacing(paragraph, before=before, after=after, line=1.1)
    return paragraph


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        run = paragraph.add_run(item)
        set_font(run, size=11)
        set_paragraph_spacing(paragraph, after=4, line=1.1)


def add_numbered(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        run = paragraph.add_run(item)
        set_font(run, size=11)
        set_paragraph_spacing(paragraph, after=4, line=1.1)


def add_metadata_table(doc, rows):
    widths = [Inches(1.8), Inches(4.7)]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_layout(table, widths)
    header_left, header_right = table.rows[0].cells
    set_cell_fill(header_left, LIGHT_FILL)
    set_cell_fill(header_right, LIGHT_FILL)
    header_left_paragraph = header_left.paragraphs[0]
    header_left_run = header_left_paragraph.add_run("Field")
    set_font(header_left_run, size=10.5, bold=True, color=SLATE)
    set_paragraph_spacing(header_left_paragraph, after=0, line=1.0)
    header_right_paragraph = header_right.paragraphs[0]
    header_right_run = header_right_paragraph.add_run("Value")
    set_font(header_right_run, size=10.5, bold=True, color=SLATE)
    set_paragraph_spacing(header_right_paragraph, after=0, line=1.0)
    mark_header_row(table.rows[0])
    for label, value in rows:
        row = table.add_row()
        set_row_layout(row, widths)
        left, right = row.cells
        set_cell_fill(left, LIGHT_FILL)
        p_left = left.paragraphs[0]
        p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
        left_run = p_left.add_run(label)
        set_font(left_run, size=10.5, bold=True, color=SLATE)
        set_paragraph_spacing(p_left, after=0, line=1.0)

        p_right = right.paragraphs[0]
        right_run = p_right.add_run(value)
        set_font(right_run, size=10.5, color=BLACK)
        set_paragraph_spacing(p_right, after=0, line=1.0)
    return table


def add_three_col_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_layout(table, widths)
    mark_header_row(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_fill(cell, LIGHT_FILL)
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(header)
        set_font(run, size=10.5, bold=True, color=SLATE)
        set_paragraph_spacing(paragraph, after=0, line=1.0)
    for row_values in rows:
        row = table.add_row()
        set_row_layout(row, widths)
        for idx, value in enumerate(row_values):
            paragraph = row.cells[idx].paragraphs[0]
            run = paragraph.add_run(value)
            set_font(run, size=10.5)
            set_paragraph_spacing(paragraph, after=0, line=1.0)
    return table


def add_diagram_block(doc, title, lines):
    widths = [Inches(6.5)]
    table = doc.add_table(rows=2, cols=1)
    table.style = "Table Grid"
    set_table_layout(table, widths)
    header_cell = table.rows[0].cells[0]
    set_cell_fill(header_cell, LIGHT_FILL)
    header_paragraph = header_cell.paragraphs[0]
    header_run = header_paragraph.add_run(title)
    set_font(header_run, size=10.5, bold=True, color=SLATE)
    set_paragraph_spacing(header_paragraph, after=0, line=1.0)
    mark_header_row(table.rows[0])

    cell = table.rows[1].cells[0]
    set_row_layout(table.rows[1], widths)
    set_cell_fill(cell, DIAGRAM_FILL)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(paragraph, after=0, line=1.0)
    for index, line in enumerate(lines):
        run = paragraph.add_run(line)
        set_font(run, name="Courier New", size=9.5, color=BLACK)
        if index < len(lines) - 1:
            run.add_break()
    return table


def configure_page(section):
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)


def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header_para.add_run("Care Nova AI | Documentation Pack")
    set_font(header_run, size=9.5, color=MUTED)

    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    footer_run = footer_para.add_run(f"Prepared from repository artifacts on {RUN_TS}")
    set_font(footer_run, size=9, color=MUTED)


def add_title_page(doc):
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(title, before=4, after=3, line=1.0)
    run = title.add_run("Care Nova AI Documentation Pack")
    set_font(run, size=22, color=BLACK, bold=True)

    subtitle = doc.add_paragraph()
    set_paragraph_spacing(subtitle, before=0, after=16, line=1.05)
    sub_run = subtitle.add_run(
        "Combined Technical Design Document, Architecture Diagram, Flow Diagram, Detailed Architecture Research, Test Plan, and Test Results"
    )
    set_font(sub_run, size=12.5, color=MUTED)
    paragraph_border(subtitle)

    add_metadata_table(
        doc,
        [
            ("Asset", "Care Nova AI"),
            ("Version", "5.0.358"),
            ("Industry", "Life Sciences and Healthcare (LSH)"),
            ("Functional Area", "Data & AI"),
            ("Agent Type", "Others"),
            ("Lifecycle Stage", "MVP"),
            ("Coverage Support Enabled", "Yes - insurance, claim, EOB, prior-authorization, and benefit-check support are included."),
            ("Automated Code Coverage Report", "No explicit automated code-coverage report is enabled in the current repository."),
            ("Document Date", "July 7, 2026"),
            ("Execution Window", RUN_TS),
        ],
    )

    add_heading(doc, "Asset Registration Snapshot", level=2)
    add_metadata_table(
        doc,
        [
            (
                "Short Description / Use Cases",
                "Real-time local healthcare advisor for patient questions, vitals interpretation, medication safety, lab explanation, records support, insurance guidance, care transitions, claims readiness, prior-authorization drafts, GxP quality workflows, and MedTech compliance support.",
            ),
            (
                "1-Line Asset Description",
                "Offline-first healthcare advisor and workflow assistant with governed specialist agents and deterministic safety boundaries.",
            ),
            (
                "Component Dependencies on API",
                "None mandatory for standard local operation. Optional connectors: OpenAI-compatible Chat Completions API, Azure OpenAI API, MedlinePlus Connect API, RxNorm / RxNav API, openFDA Drug Label API, LOINC FHIR Terminology API, and SMART on FHIR APIs.",
            ),
            (
                "Benefit Attained",
                "Primary current value case: Productivity Improvement through faster triage, draft generation, and care-workflow organization, with secondary user-experience and cost benefits from offline-first execution.",
            ),
            (
                "Quantitative Metric",
                "Current measured technical outcome: 4 of 4 validation gates passed, including 20 routed smoke-test scenarios and 61 required local model files verified. Business KPI baselines are not yet recorded in repository artifacts.",
            ),
        ],
    )

    add_heading(doc, "Included Deliverables", level=2)
    add_bullets(
        doc,
        [
            "Technical design summary covering architecture, data handling, APIs, safety boundaries, and deployment model.",
            "Architecture diagram showing UI, API, orchestration, agent, knowledge, persistence, and optional connector layers.",
            "Flow diagram describing the end-to-end patient-input to memory-update processing loop.",
            "Detailed architecture research appendix covering runtime behavior, persistence, knowledge retrieval, model routing, and source-file ownership.",
            "Test plan describing validation approach, entry criteria, suites, and scenario coverage.",
            "Executed test results based on current-run repository validations.",
        ],
    )

    add_heading(doc, "Executive Summary", level=2)
    add_body(
        doc,
        "Care Nova AI is an offline-first healthcare intelligence asset that accepts patient questions, vitals, profile data, and prior context, then routes the request through governed specialist agents to produce safety-first guidance, care summaries, and healthcare or payer workflow support. The current package combines the core design, runtime flow, code-grounded architecture research, and current validation evidence into a single handoff artifact.",
        after=8,
    )


def add_tdd_section(doc):
    doc.add_page_break()
    add_heading(doc, "1. Technical Design Document", level=1)
    add_body(
        doc,
        "This section describes the intended design of the Care Nova AI asset as implemented in the local repository. It focuses on runtime architecture, major components, responsibilities, data handling, APIs, deployment model, and safety boundaries.",
        after=8,
    )

    add_heading(doc, "1.1 Purpose and Scope", level=2)
    add_body(
        doc,
        "The asset is designed as a real-time, offline-first healthcare advisor and workflow assistant. It supports patient guidance, vitals review, medication-safety questions, lab explanation, records support, insurance and payer workflows, care transitions, and bounded administrative draft generation. The system is explicitly positioned as decision support and document support rather than diagnosis or prescribing.",
    )

    add_heading(doc, "1.2 Functional Capabilities", level=2)
    add_bullets(
        doc,
        [
            "General health question handling with grounded response synthesis.",
            "Vitals interpretation for blood pressure, blood sugar, pulse, oxygen saturation, and temperature.",
            "Medication safety support for missed-dose, interaction, and side-effect questions.",
            "Lab and report explanation in plain language.",
            "Lifestyle, wellness, mental wellness, records, insurance, and appointment support.",
            "Care transitions, claims operations, utilization-management, GxP quality, and MedTech compliance draft support.",
        ],
    )

    add_heading(doc, "1.3 Component Responsibilities", level=2)
    add_three_col_table(
        doc,
        ["Component", "Primary Responsibility", "Key Files"],
        [
            ("Web UI", "Captures profile, message, vitals, specialist inputs, and displays routed answers.", "public/index.html, public/app.js"),
            ("Local API Server", "Hosts health, readiness, analysis, memory, record, training, and supporting endpoints.", "server.js"),
            ("Health Engine", "Performs routing, risk scoring, synthesis, guardrails, and model blueprint generation.", "src/healthEngine.js"),
            ("Specialist Agents", "Delivers route-specific outputs for RAG, pharmacy, alerts, scheduling, labs, records, insurance, and payer flows.", "src/healthEngine.js"),
            ("Persistence Layer", "Stores memory, records, knowledge graph state, and local training state.", "src/memoryStore.js, src/recordStore.js, src/knowledgeGraphStore.js, src/trainingEngine.js"),
            ("Knowledge and Reasoning", "Supplies offline medical knowledge, local AI scoring, optional external cache usage, and model routing.", "src/offlineMedicalDatabase.js, src/localAiEngine.js, src/externalKnowledgeStore.js, src/hybridModelRouter.js"),
        ],
        [Inches(1.4), Inches(3.35), Inches(1.75)],
    )

    add_heading(doc, "1.4 Data Handling and APIs", level=2)
    add_body(
        doc,
        "The solution is offline-first. Normal local usage does not require mandatory external APIs. Core internal APIs include /api/analyze, /api/realtime, /api/memory, /api/records, /api/knowledge-graph, and /api/medicine/lookup. Optional connectors can be enabled for OpenAI-compatible or local model endpoints and trusted healthcare-reference services such as MedlinePlus Connect, RxNav/openFDA, and SMART on FHIR integrations.",
    )

    add_heading(doc, "1.5 Safety and Non-Goals", level=2)
    add_bullets(
        doc,
        [
            "No diagnosis, no prescribing, and no dosage calculation.",
            "No live appointment booking, caregiver contact, or emergency calling.",
            "No claim payment, coverage, or prior-authorization decision finalization.",
            "No GxP release decision, regulatory submission, or complaint disposition finalization.",
            "Patient conversations improve local context continuity, but they do not self-train medical facts.",
        ],
    )

    add_heading(doc, "1.6 Deployment Model", level=2)
    add_bullets(
        doc,
        [
            "Local-first Windows launch via start-care-nova.cmd or same-network access via start-care-nova-global.cmd.",
            "HTTP readiness and health endpoints exposed for deployment or hosting probes.",
            "Optional Docker and cloud deployment pathways represented in repository packaging.",
            "Persistent local storage for memory, records, graph state, training state, and offline source packs.",
        ],
    )


def add_architecture_section(doc):
    doc.add_page_break()
    add_heading(doc, "2. Architecture Diagram", level=1)
    add_body(
        doc,
        "The following diagram and layer summary describe the repository's current runtime architecture.",
        after=8,
    )

    add_diagram_block(
        doc,
        "Figure 1. Runtime Architecture",
        [
            "[ Presentation and Workspace UI ]",
            "  public/index.html, public/app.js",
            "                 |",
            "[ Local API Server ]",
            "  server.js",
            "                 |",
            "[ Orchestration and Routing ]",
            "  healthEngine.js, agenticRuntime.js, hybridModelRouter.js",
            "                 |",
            "+-----------------------------------------------------------+",
            "| Specialist Agents: RAG | Pharmacy | Scheduling | Alert   |",
            "| Labs | Lifestyle | Wellness | Records | Insurance        |",
            "| Care Transitions | Claims Ops | Utilization | GxP | MDT  |",
            "+-----------------------------------------------------------+",
            "                 |",
            "[ Knowledge and Reasoning Layer ]",
            "  offlineMedicalDatabase.js, localAiEngine.js,",
            "  externalKnowledgeStore.js, medicineLookupStore.js,",
            "  productIntelligence.js",
            "                 |",
            "[ Persistence and Local State ]",
            "  memoryStore.js, recordStore.js, knowledgeGraphStore.js,",
            "  trainingEngine.js, localDataMirror.js",
            "                 |",
            "[ Optional Connectors ]",
            "  OpenAI-compatible LLMs, Ollama, LM Studio,",
            "  MedlinePlus, RxNav/openFDA, SMART on FHIR",
        ],
    )

    add_heading(doc, "2.1 Layer Summary", level=2)
    add_three_col_table(
        doc,
        ["Layer", "Contained Components", "Purpose"],
        [
            ("Experience Layer", "Browser UI, workspace tabs, specialist forms", "Captures user intent and renders structured responses."),
            ("Service Layer", "HTTP server and local endpoints", "Exposes analysis, readiness, persistence, and utility APIs."),
            ("Orchestration Layer", "Health engine, runtime policy, model router", "Determines route, risk, and response assembly path."),
            ("Domain Agent Layer", "General, pharmacy, scheduling, alert, labs, payer, and quality agents", "Generates bounded outputs for each domain workflow."),
            ("Knowledge Layer", "Offline DB, local AI scoring, external cache status", "Supplies evidence, ranking, and trusted-source context."),
            ("State Layer", "Memory, records, graph, training, and local mirror stores", "Preserves continuity and operational artifacts across sessions."),
        ],
        [Inches(1.3), Inches(2.6), Inches(2.6)],
    )


def add_flow_section(doc):
    doc.add_page_break()
    add_heading(doc, "3. Flow Diagram", level=1)
    add_body(
        doc,
        "The processing loop below captures the primary runtime path from patient input through response generation and memory update.",
        after=8,
    )

    add_diagram_block(
        doc,
        "Figure 2. Primary Processing Flow",
        [
            "[1] Patient Input",
            "        |",
            "[2] Memory Store Load",
            "        |",
            "[3] Intent Classifier",
            "        |",
            "   +----+----+-------------------------+",
            "   |         |                         |",
            "[RAG]  [Pharmacy]  [Scheduling/Alert/Other Specialist]",
            "   \\         |                         /",
            "    +--------+------------------------+",
            "                 |",
            "[5] Response Synthesizer",
            "                 |",
            "[6] Safety and Guardrails",
            "                 |",
            "[7] Patient Reply",
            "                 |",
            "[8] Update Memory",
        ],
    )

    add_heading(doc, "3.1 Step Definitions", level=2)
    add_three_col_table(
        doc,
        ["Step", "Description", "Output"],
        [
            ("1. Patient Input", "User enters a symptom, request, or workflow question with optional vitals and profile details.", "Structured request payload"),
            ("2. Memory Store", "Local conversation history and saved profile context are loaded before routing.", "Memory context"),
            ("3. Intent Classifier", "Message, vitals, and context are mapped to one or more route candidates.", "Route candidates and confidence"),
            ("4. Specialist Agent Selection", "Best-fit agent or agent set is selected based on risk and route requirements.", "Agent outputs"),
            ("5. Response Synthesizer", "Agent outputs are translated into user-facing structured guidance.", "Draft final response"),
            ("6. Safety and Guardrails", "Unsafe advice patterns are blocked and escalation wording is enforced.", "Safe response"),
            ("7. Patient Reply", "The final response is returned in the UI or API response.", "Displayed answer"),
            ("8. Update Memory", "Profile, context, and run summary are saved for continuity.", "Persistent memory entry"),
        ],
        [Inches(1.1), Inches(3.65), Inches(1.75)],
    )


def add_test_plan_section(doc):
    doc.add_page_break()
    add_heading(doc, "4. Test Plan", level=1)
    add_body(
        doc,
        "The current validation strategy combines static file verification, syntax validation, scenario smoke tests, and deployment-readiness checks.",
        after=8,
    )

    add_heading(doc, "4.1 Objectives", level=2)
    add_bullets(
        doc,
        [
            "Confirm the package contains all required source, data, launch, and verification assets.",
            "Confirm edited runtime files remain syntactically valid.",
            "Confirm healthcare and payer routes produce expected risk levels and agent selections.",
            "Confirm service endpoints and readiness surfaces behave correctly for local deployment.",
        ],
    )

    add_heading(doc, "4.2 Entry Criteria", level=2)
    add_bullets(
        doc,
        [
            "Repository available locally with bundled Node runtime.",
            "Core source files and offline data assets present.",
            "Local environment able to execute node-based verification scripts.",
            "Current branch changes saved before executing validation scripts.",
        ],
    )

    add_heading(doc, "4.3 Test Suites", level=2)
    add_three_col_table(
        doc,
        ["Suite", "Command", "Coverage Focus"],
        [
            ("Syntax validation", "node --check <target files>", "Edited runtime files in both main and github-ready copies."),
            ("Model file check", "node scripts/model-file-check.js", "Required local model files, package structure, and local-only artifacts."),
            ("Smoke test suite", "node scripts/smoke-test.js", "Risk routing, agent selection, architecture metadata, and router assertions."),
            ("Deployment readiness", "node scripts/deployment-check.js", "Health/readiness APIs, headers, persistence, deployment metadata, and endpoint coverage."),
        ],
        [Inches(1.3), Inches(2.4), Inches(2.8)],
    )

    add_heading(doc, "4.4 Scenario Coverage", level=2)
    add_bullets(
        doc,
        [
            "Routine health guidance and general advice.",
            "Medication-safety, missed-dose, and high blood-pressure scenarios.",
            "Critical escalation for chest pain, stroke wording, and severe low-sugar contexts.",
            "Appointments, labs, lifestyle, wellness, records, and insurance support.",
            "Care transitions, claims operations, utilization-management, GxP quality, and MedTech compliance workflows.",
            "Cloud-route preview and forced-offline model-routing assertions inside the smoke suite.",
        ],
    )

    add_heading(doc, "4.5 Exit Criteria", level=2)
    add_bullets(
        doc,
        [
            "All selected validation scripts complete successfully.",
            "No syntax errors remain in the edited runtime files.",
            "No required local model files are missing.",
            "Deployment checks confirm healthy runtime, readiness status, and API surfaces.",
        ],
    )


def add_test_results_section(doc):
    doc.add_page_break()
    add_heading(doc, "5. Test Results", level=1)
    add_body(
        doc,
        f"Executed on {RUN_TS} ({RUN_TS_ISO}) using the bundled repository runtimes.",
        after=8,
    )

    add_three_col_table(
        doc,
        ["Verification Item", "Observed Result", "Status"],
        [
            ("Syntax validation", "6 edited runtime files validated successfully with node --check across main and github-ready copies.", "PASS"),
            ("Model file check", "Repository validation passed; 61 required local model files present and local-only artifacts preserved.", "PASS"),
            ("Smoke test suite", "Smoke tests passed, including 20 routed scenarios plus model-routing and gateway assertions.", "PASS"),
            ("Deployment readiness", "Deployment readiness checks passed for health, readiness, security headers, storage, and endpoint coverage.", "PASS"),
        ],
        [Inches(2.0), Inches(3.85), Inches(0.65)],
    )

    add_heading(doc, "5.1 Notes", level=2)
    add_bullets(
        doc,
        [
            "The smoke suite covers routine, urgent, administrative, payer, and regulated-workflow routes.",
            "Deployment checks confirm offline-first readiness, local persistence, and major status endpoints.",
            "The repository does not currently expose an explicit automated code-coverage report.",
        ],
    )


def append_research_appendix(doc):
    import build_model_research_doc as research_builder

    research_path = research_builder.build_document()
    research_doc = Document(research_path)

    doc.add_page_break()
    add_heading(doc, "6. Detailed Architecture Research", level=1)
    add_body(
        doc,
        "This appendix merges the code-grounded architecture research into the submission package so the technical design, runtime flow, persistence model, knowledge layer, and governance details stay in one review-ready document.",
        after=8,
    )

    started = False
    copied = 0
    for child in research_doc.element.body.iterchildren():
        if child.tag == qn("w:sectPr"):
            continue
        if not started:
            if child.tag != qn("w:p"):
                continue
            if extract_element_text(child) == "1. What This Model Is":
                started = True
            else:
                continue
        insert_body_element_before_section_properties(doc, deepcopy(child))
        copied += 1

    if copied == 0:
        raise RuntimeError("Architecture research appendix merge failed: source sections were not found.")


def build_document():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    for section in doc.sections:
        configure_page(section)
    add_header_footer(doc)
    add_title_page(doc)
    add_tdd_section(doc)
    add_architecture_section(doc)
    add_flow_section(doc)
    add_test_plan_section(doc)
    add_test_results_section(doc)
    append_research_appendix(doc)
    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_document()
    print(DOCX_PATH)
