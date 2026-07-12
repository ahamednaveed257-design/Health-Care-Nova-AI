from __future__ import annotations

import json
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "outputs" / "documentation-pack"
DOCX_PATH = OUTPUT_DIR / "Care-Nova-AI-Architecture-Research.docx"

TITLE_BLUE = RGBColor(0x14, 0x3D, 0x6B)
ACCENT_BLUE = RGBColor(0x2E, 0x74, 0xB5)
DEEP_BLUE = RGBColor(0x1F, 0x4D, 0x78)
TEXT_DARK = RGBColor(0x11, 0x21, 0x33)
TEXT_MUTED = RGBColor(0x57, 0x67, 0x79)
RULE = "D9E1EA"
LIGHT_FILL = "F3F6FA"
SOFT_FILL = "F8FAFC"


def load_json(path: Path):
    return json.loads(path.read_text(encoding="utf8"))


def set_run_font(run, name="Calibri", size=11, color=TEXT_DARK, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def set_paragraph_spacing(paragraph, before=0, after=0, line=1.1):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        element = tc_mar.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_layout(table, widths):
    table.autofit = False
    for idx, width in enumerate(widths):
        table.columns[idx].width = width
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
            set_cell_margins(row.cells[idx])


def mark_header_row(row):
    tr_pr = row._tr.find(qn("w:trPr"))
    if tr_pr is None:
        tr_pr = OxmlElement("w:trPr")
        row._tr.insert(0, tr_pr)
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.add_run("Care Nova AI | Architecture Research")
    set_run_font(header_run, size=9.5, color=TEXT_MUTED)
    set_paragraph_spacing(header, after=0, line=1.0)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("Page ")
    set_run_font(footer_run, size=9.5, color=TEXT_MUTED)
    add_page_number(footer)
    set_paragraph_spacing(footer, after=0, line=1.0)


def add_title(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, size=28, color=TITLE_BLUE, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, before=0, after=4, line=1.0)
    return p


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, size=13, color=TEXT_MUTED)
    set_paragraph_spacing(p, before=0, after=12, line=1.15)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    r = p.add_run(text)
    if level == 1:
        set_run_font(r, size=16, color=ACCENT_BLUE, bold=True)
        set_paragraph_spacing(p, before=16, after=8, line=1.1)
    elif level == 2:
        set_run_font(r, size=13, color=ACCENT_BLUE, bold=True)
        set_paragraph_spacing(p, before=12, after=6, line=1.1)
    else:
        set_run_font(r, size=11.5, color=DEEP_BLUE, bold=True)
        set_paragraph_spacing(p, before=8, after=4, line=1.1)
    return p


def add_body(doc, text, before=0, after=6, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, size=11, color=TEXT_DARK, italic=italic)
    set_paragraph_spacing(p, before=before, after=after, line=1.15)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run_font(r, size=11, color=TEXT_DARK)
        set_paragraph_spacing(p, before=0, after=4, line=1.15)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_run_font(r, size=11, color=TEXT_DARK)
        set_paragraph_spacing(p, before=0, after=4, line=1.15)


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_layout(table, [Inches(6.5)])
    cell = table.rows[0].cells[0]
    set_cell_fill(cell, LIGHT_FILL)
    p = cell.paragraphs[0]
    title_run = p.add_run(f"{title}: ")
    set_run_font(title_run, size=10.5, color=DEEP_BLUE, bold=True)
    body_run = p.add_run(body)
    set_run_font(body_run, size=10.5, color=TEXT_DARK)
    set_paragraph_spacing(p, before=0, after=0, line=1.1)
    doc.add_paragraph()


def add_key_value_table(doc, rows, header=("Field", "Value"), widths=(Inches(2.0), Inches(4.5))):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_layout(table, list(widths))
    mark_header_row(table.rows[0])
    for idx, text in enumerate(header):
        cell = table.rows[0].cells[idx]
        set_cell_fill(cell, LIGHT_FILL)
        p = cell.paragraphs[0]
        r = p.add_run(text)
        set_run_font(r, size=10.5, color=TEXT_MUTED, bold=True)
        set_paragraph_spacing(p, after=0, line=1.0)

    for left_text, right_text in rows:
        row = table.add_row()
        for idx, value in enumerate([left_text, right_text]):
            p = row.cells[idx].paragraphs[0]
            r = p.add_run(value)
            set_run_font(r, size=10.5)
            set_paragraph_spacing(p, after=0, line=1.0)

    doc.add_paragraph()
    return table


def add_three_col_table(doc, header, rows, widths):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_layout(table, list(widths))
    mark_header_row(table.rows[0])
    for idx, text in enumerate(header):
        cell = table.rows[0].cells[idx]
        set_cell_fill(cell, LIGHT_FILL)
        p = cell.paragraphs[0]
        r = p.add_run(text)
        set_run_font(r, size=10.25, color=TEXT_MUTED, bold=True)
        set_paragraph_spacing(p, after=0, line=1.0)

    for values in rows:
        row = table.add_row()
        for idx, value in enumerate(values):
            p = row.cells[idx].paragraphs[0]
            r = p.add_run(value)
            set_run_font(r, size=10.0, color=TEXT_DARK)
            set_paragraph_spacing(p, after=0, line=1.0)

    doc.add_paragraph()
    return table


def add_code_block(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_layout(table, [Inches(6.5)])
    cell = table.rows[0].cells[0]
    set_cell_fill(cell, SOFT_FILL)
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, before=0, after=0, line=1.0)
    for index, line in enumerate(lines):
        r = p.add_run(line)
        set_run_font(r, name="Courier New", size=9.5, color=TEXT_DARK)
        if index < len(lines) - 1:
            r.add_break()
    doc.add_paragraph()
    return table


def parse_api_endpoints(server_text: str):
    endpoints = sorted(set(re.findall(r'requestUrl\.pathname === "([^"]+)"', server_text)))
    return [endpoint for endpoint in endpoints if endpoint.startswith("/api/")]


def categorize_endpoints(endpoints):
    groups = {
        "Runtime, health, and deployment": [],
        "Clinical analysis and report generation": [],
        "Patient persistence and mirrors": [],
        "Training and calibration": [],
        "Knowledge, governance, and integrations": []
    }

    for endpoint in endpoints:
        if endpoint in {
            "/api/health", "/api/ready", "/api/readiness", "/api/deployment", "/api/deployment-readiness",
            "/api/model", "/api/model-router", "/api/model-router/preview", "/api/local-ai", "/api/model-health",
            "/api/agentic-runtime"
        }:
            groups["Runtime, health, and deployment"].append(endpoint)
        elif endpoint in {
            "/api/analyze", "/api/realtime", "/api/safety-triage", "/api/multimodal-intake", "/api/prevention-plan",
            "/api/doctor-ready-report", "/api/evidence-citations", "/api/human-review", "/api/medicine/lookup"
        }:
            groups["Clinical analysis and report generation"].append(endpoint)
        elif endpoint.startswith("/api/memory") or endpoint.startswith("/api/records") or endpoint.startswith("/api/knowledge-graph") or endpoint.startswith("/api/local-data-mirror"):
            groups["Patient persistence and mirrors"].append(endpoint)
        elif endpoint.startswith("/api/training"):
            groups["Training and calibration"].append(endpoint)
        else:
            groups["Knowledge, governance, and integrations"].append(endpoint)

    return groups


def format_endpoint_bullets(items):
    return [", ".join(items[index:index + 4]) for index in range(0, len(items), 4)]


def build_architecture_summary():
    package_json = load_json(ROOT / "package.json")
    manifest = load_json(ROOT / "data" / "offline-repository-manifest.json")
    repository = load_json(ROOT / "data" / "offline-clinical-repository.json")
    env_text = (ROOT / ".env.example").read_text(encoding="utf8")
    server_text = (ROOT / "server.js").read_text(encoding="utf8")
    readme_text = (ROOT / "README.md").read_text(encoding="utf8")

    endpoints = parse_api_endpoints(server_text)
    endpoint_groups = categorize_endpoints(endpoints)
    src_files = sorted((ROOT / "src").glob("*.js"))
    public_files = sorted((ROOT / "public").glob("*"))
    scripts = package_json.get("scripts", {})
    local_models = re.search(r"CARE_NOVA_LOCAL_MODELS=(.+)", env_text)
    local_model_list = [item.strip() for item in (local_models.group(1).split(",") if local_models else []) if item.strip()]

    return {
        "generated_at": datetime.now().astimezone().strftime("%B %d, %Y %H:%M %Z"),
        "version": package_json.get("version", "unknown"),
        "node_engine": package_json.get("engines", {}).get("node", "unknown"),
        "script_count": len(scripts),
        "scripts": scripts,
        "endpoint_count": len(endpoints),
        "endpoint_groups": endpoint_groups,
        "src_file_count": len(src_files),
        "public_file_count": len(public_files),
        "local_models": local_model_list,
        "repository_record_count": repository.get("recordCount", len(repository.get("records", []))),
        "source_pack_count": len(repository.get("sourcePacks", [])),
        "source_family_count": len(repository.get("sourceRegistry", [])),
        "manifest_summary": manifest.get("summary", {}),
        "readme_excerpt": readme_text
    }


def build_document():
    summary = build_architecture_summary()
    doc = Document()
    configure_document(doc)

    title = add_title(doc, "Care Nova AI")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_subtitle(doc, "Architecture, build logic, runtime workflow, knowledge retrieval, storage model, and operating behavior")
    add_body(
        doc,
        "This document is a code-grounded technical research note built from the current Care Nova AI repository. "
        "It explains what the model actually is, how it is built, how a request flows through the system, where data lives, "
        "how offline and online paths differ, and which files implement each major subsystem.",
        after=10
    )

    add_key_value_table(
        doc,
        [
            ("App version", summary["version"]),
            ("Node runtime target", summary["node_engine"]),
            ("Detected API endpoints", str(summary["endpoint_count"])),
            ("Source files under src/", str(summary["src_file_count"])),
            ("Offline repository records", str(summary["manifest_summary"].get("totalRetrievalRecords", summary["repository_record_count"]))),
            ("Source families", str(summary["source_family_count"])),
            ("Generated from repo state at", summary["generated_at"])
        ]
    )

    add_callout(
        doc,
        "Research conclusion",
        "Care Nova AI is not a single monolithic medical LLM. It is a local-first agentic healthcare system composed of a browser UI, "
        "a Node.js API server, a deterministic orchestration engine, an offline retrieval layer, persistent local patient stores, "
        "and optional open-source or cloud LLM refinement stages."
    )

    add_heading(doc, "1. What This Model Is", level=1)
    add_body(
        doc,
        "At runtime, Care Nova AI behaves as a multi-stage orchestration system rather than as a generic chatbot. "
        "The core of the product is the pairing of server.js and src/healthEngine.js. The server owns request handling, "
        "health probes, data loading, persistence, and post-processing. The health engine owns normalization, intent routing, "
        "risk scoring, knowledge retrieval, agent execution, synthesis, and guardrails."
    )
    add_bullets(
        doc,
        [
            "Local-first healthcare workspace: the default path is designed to work with local storage, local knowledge, and local fallbacks before any online dependency is considered.",
            "Agentic workflow engine: user requests are classified, routed, executed through one or more specialist agents, and consolidated into a single guarded answer.",
            "Retrieval-augmented and memory-aware: the system combines offline medical records, patient memory, saved records, and a patient knowledge graph when reasoning.",
            "LLM-optional architecture: deterministic outputs always exist; local open-source and cloud models are treated as optional enhancement layers.",
            "Safety-bounded runtime: the answer is explicitly constrained by no-diagnosis, no-prescribing, no-dose-calculation, and no-live-action boundaries."
        ]
    )

    add_heading(doc, "2. The Build Strategy", level=1)
    add_body(
        doc,
        "The repository is organized as a layered web application. The model is built by composing focused modules rather than by embedding all logic in a single file. "
        "That separation is the main reason the project can support many tabs, many specialist routes, offline operation, GitHub packaging, and multiple runtime modes."
    )
    add_three_col_table(
        doc,
        ("Layer", "Main files", "Responsibility"),
        [
            ("Browser shell", "public/index.html, public/app.js, public/styles.css, public/sw.js", "Renders the workspace, manages tabs and modes, captures inputs, calls APIs, stores lightweight UI state in localStorage, and supports installable PWA behavior."),
            ("HTTP server", "server.js", "Serves static assets, exposes the API surface, loads memory and records, invokes the core engine, applies optional LLM assists, and persists outputs."),
            ("Core orchestration", "src/healthEngine.js, src/agenticRuntime.js", "Normalizes requests, classifies intents, scores risk, builds plans, runs agents, synthesizes responses, and explains runtime decisions."),
            ("Knowledge and ranking", "src/offlineMedicalDatabase.js, src/localAiEngine.js, scripts/build-offline-repository.js", "Loads the offline repository and ranks evidence using lexical, entity, semantic-family, route, and safety-aware scoring."),
            ("Model routing and optional LLMs", "src/hybridModelRouter.js, src/openSourceLocalRuntime.js, src/localReasoningGateway.js, src/specialistLlmGateway.js, src/cloudLlmGateway.js", "Chooses local vs cloud processing, probes Ollama/LM Studio style runtimes, and performs route-aware refinement when configured."),
            ("Persistent data and mirrors", "src/memoryStore.js, src/recordStore.js, src/knowledgeGraphStore.js, src/localDataMirror.js, src/trainingEngine.js", "Stores patient memory, records, graph facts, data mirrors, and training/calibration state on the local machine.")
        ],
        widths=(Inches(1.35), Inches(2.25), Inches(2.9))
    )

    add_heading(doc, "3. End-to-End Runtime Flow", level=1)
    add_body(
        doc,
        "The main request path is implemented across POST /api/analyze in server.js and analyzeHealthQuery() in src/healthEngine.js. "
        "The sequence below reflects the current implementation, including the optional local and cloud post-processors."
    )
    add_code_block(
        doc,
        [
            "Browser UI (public/app.js)",
            "  -> POST /api/analyze or POST /api/realtime",
            "  -> server.js loads memory, records, external cache, training calibration",
            "  -> analyzeHealthQuery() normalizes input and builds requirement profile",
            "  -> classifyIntents() + calculateRisk()",
            "  -> retrieveMedicalKnowledge() using offline repository and optional cached external knowledge",
            "  -> buildExecutionPlan() + applySingleAgentScope() + precision supervision",
            "  -> selectHybridModelRoute() + buildLlmCognitiveCore()",
            "  -> runAgentPlan() executes route-specific deterministic specialists",
            "  -> tryEnhanceSpecialistAgentResultsWithLlm() optionally refines specialist outputs",
            "  -> synthesizeResponse() creates the patient-facing answer",
            "  -> applyGuardrails() checks no-diagnosis / no-prescribing / no-dosage wording",
            "  -> server.js optionally applies local reasoning assist and cloud second pass",
            "  -> appendPatientMemory() + upsertPatientKnowledgeGraph() + syncLocalDataMirror()",
            "  -> JSON result returned to UI and rendered in the active tab"
        ]
    )
    add_numbered(
        doc,
        [
            "The UI captures a free-text health request plus optional structured profile, vitals, and context values.",
            "The server merges the request with saved patient profile data, prior memory, stored records, and training calibration before analysis begins.",
            "The engine extracts additional vitals from the message itself, normalizes the profile and context, and builds input-quality and requirement-fit summaries.",
            "Intent classification determines which routes are relevant. Risk scoring then evaluates urgency using message content, vitals, context, and prior history.",
            "The retrieval layer builds a route-aware knowledge query and ranks local medical references from the offline repository.",
            "The planner decides whether this is a single-agent or multi-agent response and which route should own the answer.",
            "Deterministic specialist runners execute route-specific logic for general guidance, specialist review, vitals, medicines, labs, scheduling, safety, records, insurance, and operations paths.",
            "If configured, specialist LLM refinement runs on top of those deterministic outputs. This is a second pass, not the primary decision engine.",
            "The synthesizer creates one focused, patient-friendly response with actions, warning signs, support sections, and a disclaimer.",
            "Guardrails scan the final wording for blocked patterns such as diagnosis claims, prescribing language, or dosage instructions.",
            "After the answer is ready, the server can add local reasoning assist, cloud rewrite assist, evidence packets, doctor-ready notes, prevention plans, safety triage, and human review packets.",
            "The run is persisted locally into memory, records, graph facts, and mirror files so the next request starts with richer context."
        ]
    )

    add_heading(doc, "4. Agent Inventory", level=1)
    add_body(
        doc,
        "The execution engine maps routes to explicit runner functions. The patient-facing tabs are only one surface; underneath them, the engine owns a broader route catalog."
    )
    add_three_col_table(
        doc,
        ("Route", "Typical UI surface", "Purpose"),
        [
            ("RAG_AGENT", "General", "First-layer general health guidance, context gathering, warning-sign review, and next-step framing."),
            ("SPECIALIST_DOCTOR_AGENT", "Specialist", "Disease-focused structured review with domain-specific expectations, precautions, and follow-up questions."),
            ("VITALS_AGENT", "Vitals", "Interprets BP, glucose, pulse, temperature, BMI, and home-monitoring context."),
            ("PHARMACY_AGENT", "Medicine", "Medication label review, missed-dose safety, interaction questions, side-effect triage, and storage guidance."),
            ("LABS_AGENT", "Labs", "Explains pasted or uploaded lab text, extracts values, highlights abnormal items, and builds doctor questions."),
            ("SCHEDULING_AGENT", "Visits", "Follow-up planning, appointment preparation, visit type inference, readiness gaps, and communication scripts."),
            ("ALERT_AGENT", "Safety", "Urgent-signal handling, escalation wording, monitoring checklist, and handoff packet generation."),
            ("WELLNESS_AGENT / LIFESTYLE_AGENT", "Wellness", "Routine-building, stress and sleep guidance, daily health habits, and supportive coaching within safety boundaries."),
            ("RECORDS_AGENT", "Records", "Local record organization, summaries, doctor-ready packets, and timeline support."),
            ("INSURANCE_AGENT", "Insurance", "Claim organization, document gaps, benefit questions, appeal preparation, and administrative support."),
            ("CARE_TRANSITIONS_AGENT", "Transitions", "Discharge and follow-up workflow support."),
            ("CLAIMS_OPS_AGENT / UTILIZATION_AGENT / GXP_QUALITY_AGENT / MEDTECH_COMPLIANCE_AGENT", "Operations / regulated workflows", "Administrative and regulated-domain drafting helpers for broader healthcare workflow use cases.")
        ],
        widths=(Inches(1.55), Inches(1.55), Inches(3.4))
    )

    add_heading(doc, "5. The Knowledge and RAG Layer", level=1)
    add_body(
        doc,
        "Care Nova AI has a real offline knowledge subsystem. It is not a separate vector database service; instead, the repository and index are stored as JSON files and searched by a local ranking engine."
    )
    add_key_value_table(
        doc,
        [
            ("Repository file", "data/offline-clinical-repository.json"),
            ("Base seed file", "data/offline-medical-db.json"),
            ("Index file", "data/offline-knowledge-index.json"),
            ("Manifest file", "data/offline-repository-manifest.json"),
            ("Total retrieval records", str(summary["manifest_summary"].get("totalRetrievalRecords", 0))),
            ("Base seeded records", str(summary["manifest_summary"].get("baseRecordCount", 0))),
            ("Repository-generated records", str(summary["manifest_summary"].get("repositoryRecordCount", 0))),
            ("Source-pack records", str(summary["manifest_summary"].get("packRecordCount", 0))),
            ("Source pack count", str(summary["source_pack_count"])),
            ("Source family count", str(summary["source_family_count"])),
            ("Indexed tokens", str(summary["manifest_summary"].get("tokenCount", 0)))
        ]
    )
    add_body(
        doc,
        "The offline repository is built by scripts/build-offline-repository.js. That build script normalizes base records, merges modular source packs, enriches each record, "
        "builds a lexical/vector-style index, and emits a manifest containing source governance, licensing notes, maintenance workflow, and distribution counts."
    )
    add_bullets(
        doc,
        [
            "rankLocalMedicalKnowledge() in src/localAiEngine.js is the heart of local retrieval.",
            "The ranker combines tokenization, synonym expansion, exact phrase hits, TF-IDF style vector similarity, medical entity alignment, clinical domain matching, numeric vital/lab signal matching, route-tag alignment, population context, and urgent-safety boosts.",
            "Each match receives a relevance score, matched terms, semantic families, route-tag hits, and an evidence grade such as strong, good, supporting, or weak.",
            "retrieveMedicalKnowledge() in src/healthEngine.js wraps that ranker with request-aware scopes, caching, and route-specific retrieval policies.",
            "External online knowledge, when enabled, is normalized and cached locally before reuse; it does not bypass local guardrails."
        ]
    )

    add_heading(doc, "6. Model Routing and LLM Participation", level=1)
    add_body(
        doc,
        "The project uses the word 'model' in several layers. The important distinction is that the deterministic local core always exists, while optional LLM runtimes refine or rewrite parts of the output when available."
    )
    add_three_col_table(
        doc,
        ("Layer", "Main file", "What it actually does"),
        [
            ("Hybrid router", "src/hybridModelRouter.js", "Scores request complexity, safety context, policy, and connectivity to choose local, cloud, or hybrid processing."),
            ("Local runtime probe", "src/openSourceLocalRuntime.js", "Detects and probes local OpenAI-compatible runtimes such as Ollama and LM Studio, checks reachable endpoints, and confirms model availability."),
            ("Specialist assist", "src/specialistLlmGateway.js", "Optionally refines deterministic specialist agent outputs for up to a bounded set of routes."),
            ("Local reasoning assist", "src/localReasoningGateway.js", "Runs a route-aware open-source local second pass after the main analysis and only applies changes if final guardrails still pass."),
            ("Cloud rewrite assist", "src/cloudLlmGateway.js", "Provides an opt-in cloud rewrite or route-aware second pass when environment flags explicitly allow paid/cloud usage.")
        ],
        widths=(Inches(1.25), Inches(2.0), Inches(3.25))
    )
    add_body(
        doc,
        "The default environment file shows that the project is designed to prefer local models and the lowest-cost route. "
        f"The currently declared local open-source families are: {', '.join(summary['local_models']) or 'none declared'}."
    )
    add_bullets(
        doc,
        [
            "The hybrid router can assemble an open-source participation plan with a primary reasoner, verifier, and responder when more than one local model is available.",
            "If cloud access is disabled or unavailable, the router falls back to the best local model and then to the deterministic Care Nova core.",
            "The local reasoning assist and specialist LLM assist are enhancement paths; they do not replace the baseline deterministic health engine.",
            "Because all LLM assists are wrapped around guardrails and route boundaries, the project tries to preserve safety even when a model is configured."
        ]
    )

    add_heading(doc, "7. What Machine Learning Exists Today", level=1)
    add_body(
        doc,
        "The repository includes a training engine, but it is important to describe it precisely. The current implementation is a lightweight calibration system, not a full local medical foundation-model training stack."
    )
    add_bullets(
        doc,
        [
            "src/trainingEngine.js stores approved examples in a local JSON training file and builds route priors, keyword-route weights, route reliability scores, and a confusion matrix.",
            "The training metadata explicitly reports noPhiTraining: true and medicalFactTraining: false. That means patient examples are intended for routing and calibration, not for learning new medical facts.",
            "The calibration model improves route fit and reliability scoring; it does not turn the app into a self-updating clinical knowledge source.",
            "The retrieval layer itself behaves like a compact local ML ranker through weighted scoring and TF-IDF style vectors, but it is still deterministic and repository-bounded."
        ]
    )
    add_callout(
        doc,
        "Accuracy note",
        "The strongest accuracy path in this codebase comes from the combination of route control, risk gates, offline knowledge retrieval, and structured persistence. "
        "The optional LLM layers improve phrasing or specialist refinement, but they are not the only source of correctness."
    )

    add_heading(doc, "8. Persistence, Privacy, and Local Data", level=1)
    add_body(
        doc,
        "Patient context is stored on the local machine rather than in a remote database by default. The project separates several types of data so that the next interaction can reuse context without collapsing everything into one flat file."
    )
    add_three_col_table(
        doc,
        ("Store", "Path", "Role"),
        [
            ("Memory store", "data/memory/patient-memory.json", "Conversation history and memory patches. Capped history is merged and persisted per patient."),
            ("Record store", "data/records/patient-records.json", "Patient records, summaries, and selected record metadata."),
            ("Knowledge graph", "data/graph/patient-knowledge-graph.json and data/graph/patients/", "Structured graph facts, nodes, edges, sharded patient storage, and bundled graph support."),
            ("Training state", "data/training/agent-training-state.json", "Approved examples and route-calibration data."),
            ("External cache", "data/external/external-knowledge-cache.json", "Optional normalized online reference cache for later offline reuse."),
            ("OneDrive mirror", "data/onedrive-mirror/", "Local mirror copy of memory, records, graph, and related files for additional persistence in a OneDrive path.")
        ],
        widths=(Inches(1.2), Inches(2.35), Inches(2.95))
    )
    add_bullets(
        doc,
        [
            "The memory and record stores use a queue-based atomic write pattern: write to a temporary file first, then replace the target file.",
            "The knowledge graph store supports both a legacy single-file graph and sharded patient graph files, plus bundled patient shards for packaged distribution.",
            "Browser localStorage is used only for light UI state such as selected theme, interface mode, and lightweight history helpers; the primary patient memory path is the local server store.",
            "LocalDataMirror synchronizes selected local data files to a OneDrive-compatible mirror directory after key operations."
        ]
    )

    add_heading(doc, "9. API Surface", level=1)
    add_body(
        doc,
        "The server exposes a broad API surface. Some endpoints drive the main app directly, while others expose diagnostics, governance, training, or export support."
    )
    add_key_value_table(
        doc,
        [
            ("Total detected API endpoints", str(summary["endpoint_count"])),
            ("Runtime and health endpoints", str(len(summary["endpoint_groups"]["Runtime, health, and deployment"]))),
            ("Clinical action endpoints", str(len(summary["endpoint_groups"]["Clinical analysis and report generation"]))),
            ("Persistence endpoints", str(len(summary["endpoint_groups"]["Patient persistence and mirrors"]))),
            ("Training endpoints", str(len(summary["endpoint_groups"]["Training and calibration"]))),
            ("Knowledge and governance endpoints", str(len(summary["endpoint_groups"]["Knowledge, governance, and integrations"])))
        ]
    )
    for group_name, endpoints in summary["endpoint_groups"].items():
        add_heading(doc, group_name, level=2)
        add_bullets(doc, format_endpoint_bullets(endpoints))

    add_heading(doc, "10. Frontend Application Model", level=1)
    add_body(
        doc,
        "The frontend is implemented in public/app.js and public/index.html as a large vanilla JavaScript single-page application. "
        "It does not depend on a heavyweight frontend framework. That choice keeps the installation simple and makes GitHub/static packaging easier, but it also means UI state management is hand-built."
    )
    add_bullets(
        doc,
        [
            "Interface routing is handled through hash and history state helpers, allowing direct tab navigation such as General, Specialist, Atlas, Vitals, Medicine, Labs, Wellness, Visits, Safety, Records, Insurance, Summary, Profile, and Guide.",
            "The UI supports guided and expert modes, per-tab workspace subtabs, a compact command hub, multilingual translation, theme switching, and installed-app behavior.",
            "Realtime mode uses /api/realtime for low-latency preview while full runs use /api/analyze for persistence and post-processing.",
            "The Guide system includes both a slide guide and an optional chunked HD video loader path for larger walkthrough assets.",
            "The service worker and web manifest make the app installable as a PWA, enabling splash/loading behavior and offline shell caching."
        ]
    )

    add_heading(doc, "11. Offline vs Online Operation", level=1)
    add_body(
        doc,
        "The core design target is parity of behavior across offline and online modes. Online mode adds optional knowledge or rewrite stages, but the same local engine remains in charge."
    )
    add_three_col_table(
        doc,
        ("Mode", "What stays the same", "What can be added"),
        [
            ("Offline", "The same Node API, health engine, specialist runners, local knowledge, local storage, and guardrails remain active.", "No internet is required. The system uses only local retrieval, local persistence, and optional local LLM runtimes."),
            ("Online-ready local safe core", "The local deterministic path still runs first and still owns persistence and safety boundaries.", "Approved external reference retrieval and optional cloud rewrite paths can be enabled through environment variables."),
            ("Hybrid", "The planner, route selection, memory, records, graph updates, and guardrails remain local.", "Cloud or remote models may handle complex rewrites or second-pass reasoning when policy allows it.")
        ],
        widths=(Inches(1.1), Inches(2.6), Inches(2.8))
    )
    add_body(
        doc,
        "Connectivity policy is implemented in src/runtimeConnectivity.js. It determines whether a remote endpoint is usable for the current run, "
        "whether forced-offline mode is active, and whether only local endpoints are allowed."
    )

    add_heading(doc, "12. Safety and Governance", level=1)
    add_body(
        doc,
        "Safety is not a final banner added only in the UI. It is embedded across the repository: rule-based risk scoring, specialist route boundaries, retrieval scope control, final response guardrails, trusted source planning, and governance metadata."
    )
    add_bullets(
        doc,
        [
            "applyGuardrails() in src/healthEngine.js blocks diagnosis wording, dosage instructions, prescribing language, and certain categories of unsupported action claims.",
            "advancedCapabilityEngine.js contains urgent-signal rules for chest pain, stroke, severe allergy, fainting, dangerous bleeding, mental-health crisis, very high BP, and high-fever risk.",
            "productIntelligence.js documents trusted source connectors, offline packs, report templates, quality metrics, governance expectations, and FHIR-readiness boundaries.",
            "The offline repository build process explicitly states that copyrighted full textbooks or clinical guidelines should not be bundled without license approval.",
            "The training subsystem is intentionally restricted to approved feedback calibration and explicitly avoids using PHI as free-form training data."
        ]
    )

    add_heading(doc, "13. Build, Run, and Release Workflow", level=1)
    add_body(
        doc,
        "The repository includes a clear set of commands for local execution, syntax validation, smoke testing, offline repository rebuilds, local LLM checks, and GitHub packaging."
    )
    add_key_value_table(
        doc,
        [
            ("start", summary["scripts"].get("start", "")),
            ("check", summary["scripts"].get("check", "")),
            ("test", summary["scripts"].get("test", "")),
            ("llm:check", summary["scripts"].get("llm:check", "")),
            ("offline:build", summary["scripts"].get("offline:build", "")),
            ("deploy:check", summary["scripts"].get("deploy:check", "")),
            ("release:check", summary["scripts"].get("release:check", ""))
        ]
    )
    add_bullets(
        doc,
        [
            "scripts/smoke-test.js exercises representative healthcare cases and validates route, risk, runtime, and retrieval behavior.",
            "scripts/deployment-check.js stands up a temporary server and verifies health endpoints, readiness, router state, persistence status, and deployment-related invariants.",
            "scripts/model-file-check.js validates that key source, UI, and data files are present.",
            "scripts/build-github-standalone.js and the github-ready-care-nova-ai folder package a static-friendly variant, but the richest local behavior still depends on the Node backend."
        ]
    )

    add_heading(doc, "14. Practical Strengths and Real Constraints", level=1)
    add_three_col_table(
        doc,
        ("Area", "Strength", "Constraint / truth-in-advertising note"),
        [
            ("Routing", "Clear route ownership and specialist separation reduce generic chatbot drift.", "Classification quality still depends on prompt clarity and route cues in the user input."),
            ("Offline operation", "Core analysis, persistence, retrieval, and guardrails all work without internet.", "Offline knowledge is bounded by the curated repository present on disk."),
            ("LLM integration", "Supports local open-source models and optional cloud paths without making them mandatory.", "Configured model endpoints must actually be reachable and loaded; otherwise the system falls back."),
            ("Persistence", "Memory, records, graph, and mirror support create strong longitudinal context.", "Static-only deployments cannot reproduce the full local server persistence model by themselves."),
            ("Training", "Approved-example calibration can improve route fit over time.", "This is calibration logic, not a full medical deep-learning pipeline or autonomous medical fact learning.")
        ],
        widths=(Inches(1.15), Inches(2.45), Inches(2.9))
    )

    add_heading(doc, "15. How to Extend the Model Safely", level=1)
    add_numbered(
        doc,
        [
            "Add or update curated offline knowledge in data/offline-source-packs/, then rebuild the repository with scripts/build-offline-repository.js.",
            "Introduce a new specialist route by adding route classification signals, execution-plan support, a runner function in src/healthEngine.js, UI routing labels, and smoke-test coverage.",
            "If an LLM should refine that route, add a route boundary and schema support to src/specialistLlmGateway.js or src/localReasoningGateway.js.",
            "Update server endpoints only when a new capability truly needs an API surface, and keep persistence local by default.",
            "Add deployment or runtime flags to .env.example only when they can be backed by a safe fallback path."
        ]
    )

    add_heading(doc, "16. Key Source Files", level=1)
    add_key_value_table(
        doc,
        [
            ("server.js", "HTTP entrypoint, API surface, health/readiness endpoints, persistence orchestration, and post-processing."),
            ("src/healthEngine.js", "Main healthcare workflow engine and the most important file in the repository."),
            ("src/localAiEngine.js", "Offline ranking and evidence scoring logic."),
            ("src/hybridModelRouter.js", "Route-aware local/cloud model selection and failover chain builder."),
            ("src/openSourceLocalRuntime.js", "Local runtime discovery and endpoint probing for Ollama/LM Studio style adapters."),
            ("src/localReasoningGateway.js", "Optional open-source local second-pass reasoning on the final result."),
            ("src/specialistLlmGateway.js", "Optional route-level specialist output refinement."),
            ("src/memoryStore.js", "Persistent patient conversation memory."),
            ("src/recordStore.js", "Persistent patient records."),
            ("src/knowledgeGraphStore.js", "Persistent structured patient knowledge graph with sharding support."),
            ("src/trainingEngine.js", "Local feedback example store and route-calibration logic."),
            ("public/app.js", "Frontend controller for tabs, theme, language, guide, results, and API calls."),
            ("scripts/build-offline-repository.js", "Offline repository and index builder."),
            ("scripts/smoke-test.js", "Regression and scenario validation."),
            ("scripts/deployment-check.js", "Deployment readiness validation.")
        ]
    )

    add_heading(doc, "17. Bottom Line", level=1)
    add_body(
        doc,
        "Care Nova AI is best understood as a local-first healthcare operating model built out of specialized software layers. "
        "Its core strength is not that one single model 'knows everything'; its strength is that it structures input, uses explicit route control, "
        "grounds answers in local evidence, preserves context over time, and only then allows optional model-assisted refinement. "
        "That architecture is why the project can support many tabs, offline behavior, local privacy, and safer response shaping in one workspace."
    )

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    output = build_document()
    print(output)
